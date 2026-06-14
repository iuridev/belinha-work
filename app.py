from flask import (Flask, render_template, request, redirect,
                   url_for, session, flash, jsonify, send_file)
from functools import wraps
from datetime import datetime

import database as db
import sheets
import csv_parser
from config import SECRET_KEY, DISCIPLINAS, DISCIPLINAS_NOMES

app = Flask(__name__)
app.secret_key = SECRET_KEY


# ── Decorador de autenticação ──────────────────────────────────────────────

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'usuario' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated


# ── Login / Logout ─────────────────────────────────────────────────────────

@app.route('/', methods=['GET', 'POST'])
def login():
    if 'usuario' in session:
        return redirect(url_for('dashboard'))

    erro = None
    if request.method == 'POST':
        login_val = request.form.get('login', '').strip()
        senha = request.form.get('senha', '').strip()
        usuario = sheets.autenticar(login_val, senha)
        if usuario:
            session['usuario'] = dict(usuario)
            return redirect(url_for('dashboard'))
        erro = 'Login ou senha incorretos.'

    return render_template('login.html', erro=erro)


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


# ── Dashboard ──────────────────────────────────────────────────────────────

@app.route('/dashboard')
@login_required
def dashboard():
    ano = int(request.args.get('ano', datetime.now().year))
    stats = db.get_stats()
    anos = db.get_anos_disponiveis() or [ano]
    bimestre = request.args.get('bimestre')
    ranking = db.get_ranking_professores(ano, int(bimestre) if bimestre else None)
    return render_template('dashboard.html',
                           stats=stats, ranking=ranking,
                           ano=ano, anos=anos,
                           bimestre_sel=bimestre)


# ── Importar CSV ───────────────────────────────────────────────────────────

@app.route('/importar', methods=['GET', 'POST'])
@login_required
def importar():
    if request.method == 'POST':
        arquivo = request.files.get('arquivo')
        bimestre = request.form.get('bimestre')
        ano = request.form.get('ano')

        if not arquivo or not bimestre or not ano:
            flash('Preencha todos os campos e selecione o arquivo.', 'erro')
            return redirect(url_for('importar'))

        if not arquivo.filename.lower().endswith('.csv'):
            flash('O arquivo deve ser no formato CSV.', 'erro')
            return redirect(url_for('importar'))

        tipo_avaliacao = request.form.get('tipo_avaliacao', '').strip()
        if not tipo_avaliacao:
            flash('Selecione o tipo de avaliação.', 'erro')
            return redirect(url_for('importar'))

        try:
            conteudo = arquivo.read()
            registros = csv_parser.parse_csv(conteudo)

            if not registros:
                flash('Nenhuma turma válida encontrada no arquivo.', 'erro')
                return redirect(url_for('importar'))

            salvos = db.salvar_avaliacoes(registros, int(bimestre), int(ano), tipo_avaliacao)

            # Grava também na aba 'avaliacoes' do Google Sheets
            msg_sheets = ''
            try:
                gravados = sheets.salvar_avaliacoes_planilha(registros, int(bimestre), int(ano), tipo_avaliacao)
                msg_sheets = f' ({gravados} linhas salvas na planilha)'
            except FileNotFoundError:
                msg_sheets = ' (planilha não atualizada: credentials.json ausente)'
            except Exception as e_sh:
                msg_sheets = f' (erro na planilha: {e_sh})'

            flash(
                f'{salvos} turmas importadas com sucesso — {bimestre}° Bimestre / {ano}.{msg_sheets}',
                'sucesso'
            )
            return redirect(url_for('vinculos'))

        except Exception as e:
            flash(f'Erro ao processar arquivo: {e}', 'erro')
            return redirect(url_for('importar'))

    ano_atual = datetime.now().year
    anos = list(range(ano_atual - 1, ano_atual + 2))
    tipos = sheets.get_tipos_avaliacao()
    return render_template('importar.html', anos=anos, ano_atual=ano_atual, tipos=tipos)


# ── Vínculos professor × disciplina × turma ────────────────────────────────

@app.route('/vinculos')
@login_required
def vinculos():
    ano = int(request.args.get('ano', datetime.now().year))
    professores = sheets.get_professores()
    turmas = db.get_turmas_disponiveis()
    vinculos_lista = db.get_vinculos(ano)
    anos = db.get_anos_disponiveis() or [ano]

    return render_template('vinculos.html',
                           professores=professores,
                           turmas=turmas,
                           disciplinas=DISCIPLINAS,
                           disciplinas_nomes=DISCIPLINAS_NOMES,
                           vinculos=vinculos_lista,
                           ano=ano,
                           anos=anos)


@app.route('/vinculos/salvar-lote', methods=['POST'])
@login_required
def salvar_vinculos_lote():
    data           = request.json or {}
    professor_id   = data.get('professor_id')
    nome_professor = data.get('nome_professor')
    disciplina     = data.get('disciplina')
    turmas         = data.get('turmas', [])
    ano_letivo     = int(data.get('ano_letivo', datetime.now().year))

    salvos = 0
    erros  = []
    for turma in turmas:
        try:
            # 1. Grava na planilha (fonte de verdade)
            vid = sheets.salvar_vinculo_planilha(professor_id, nome_professor, disciplina, turma, ano_letivo)
            # 2. Atualiza cache SQLite
            db.cache_upsert_vinculo(vid, professor_id, nome_professor, disciplina, turma, ano_letivo)
            salvos += 1
        except Exception as e:
            erros.append(str(e))

    if erros:
        return jsonify({'ok': False, 'salvos': salvos, 'erros': erros}), 400
    return jsonify({'ok': True, 'salvos': salvos})


@app.route('/api/turmas-por-disciplina')
@login_required
def api_turmas_por_disciplina():
    disciplina = request.args.get('disciplina', '')
    ano        = int(request.args.get('ano', datetime.now().year))
    turmas     = db.get_turmas_por_disciplina(disciplina, ano)
    return jsonify(turmas)


@app.route('/vinculos/remover/<int:vid>', methods=['DELETE'])
@login_required
def remover_vinculo(vid):
    try:
        # 1. Remove da planilha (fonte de verdade)
        sheets.remover_vinculo_planilha(vid)
    except Exception as e:
        print(f'Aviso ao remover da planilha: {e}')
    # 2. Remove do cache SQLite
    db.cache_remover_vinculo(vid)
    return jsonify({'ok': True})


@app.route('/sincronizar', methods=['POST'])
@login_required
def sincronizar():
    resultado = _sincronizar_cache()
    return jsonify(resultado)


# ── Professores ────────────────────────────────────────────────────────────

@app.route('/professores')
@login_required
def professores():
    lista = sheets.get_professores()
    ano = int(request.args.get('ano', datetime.now().year))
    anos = db.get_anos_disponiveis() or [ano]
    return render_template('professores.html', professores=lista, ano=ano, anos=anos)


@app.route('/professores/novo', methods=['GET', 'POST'])
@login_required
def novo_professor():
    if request.method == 'POST':
        nome = request.form.get('nome', '').strip()
        cpf  = request.form.get('cpf', '').strip()

        if not nome:
            flash('O nome é obrigatório.', 'erro')
            return redirect(url_for('novo_professor'))

        try:
            sheets.adicionar_professor(nome, cpf)
            flash(f'Professor "{nome}" cadastrado com sucesso na planilha.', 'sucesso')
            return redirect(url_for('professores'))
        except FileNotFoundError as e:
            flash(str(e), 'erro')
        except Exception as e:
            flash(f'Erro ao salvar na planilha: {e}', 'erro')

    return render_template('novo_professor.html')


@app.route('/professores/inativar/<professor_id>', methods=['POST'])
@login_required
def inativar_professor(professor_id):
    try:
        sheets.inativar_professor(professor_id)
        flash('Professor inativado com sucesso.', 'sucesso')
    except Exception as e:
        flash(f'Erro ao inativar: {e}', 'erro')
    return redirect(url_for('professores'))


# ── Ficha do professor ─────────────────────────────────────────────────────

@app.route('/professor/<professor_id>')
@login_required
def ficha_professor(professor_id):
    lista = sheets.get_professores()
    professor = next((p for p in lista if str(p.get('id')) == str(professor_id)), None)

    if not professor:
        flash('Professor não encontrado.', 'erro')
        return redirect(url_for('professores'))

    ano  = int(request.args.get('ano', datetime.now().year))
    tipo = request.args.get('tipo', '')
    anos = db.get_anos_disponiveis() or [ano]
    tipos_disponiveis = db.get_tipos_disponiveis_professor(professor_id, ano)
    performance, bimestres = db.get_professor_performance(professor_id, ano, tipo or None)

    return render_template('ficha.html',
                           professor=professor,
                           performance=performance,
                           bimestres=bimestres,
                           ano=ano,
                           anos=anos,
                           tipo_sel=tipo,
                           tipos_disponiveis=tipos_disponiveis,
                           disciplinas_nomes=DISCIPLINAS_NOMES)


# ── Devolutiva de Índices Individuais ─────────────────────────────────────

@app.route('/professor/<professor_id>/devolutiva')
@login_required
def devolutiva_professor(professor_id):
    lista = sheets.get_professores()
    professor = next((p for p in lista if str(p.get('id')) == str(professor_id)), None)
    if not professor:
        flash('Professor não encontrado.', 'erro')
        return redirect(url_for('professores'))

    ano = int(request.args.get('ano', datetime.now().year))
    anos = db.get_anos_disponiveis() or [ano]
    tipos_disponiveis = db.get_tipos_disponiveis_professor(professor_id, ano)

    performance_por_tipo = {}
    for tipo in tipos_disponiveis:
        perf, bimestres = db.get_professor_performance(professor_id, ano, tipo)
        if perf:
            performance_por_tipo[tipo] = {'performance': perf, 'bimestres': bimestres}

    return render_template('devolutiva.html',
                           professor=professor,
                           ano=ano,
                           anos=anos,
                           performance_por_tipo=performance_por_tipo,
                           data_geracao=datetime.now().strftime('%d/%m/%Y'),
                           disciplinas_nomes=DISCIPLINAS_NOMES)


@app.route('/professor/<professor_id>/devolutiva/word', methods=['POST'])
@login_required
def devolutiva_word(professor_id):
    import io, base64 as _b64
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        flash('Biblioteca python-docx não instalada. Execute: pip install python-docx', 'erro')
        return redirect(url_for('devolutiva_professor', professor_id=professor_id))

    lista = sheets.get_professores()
    professor = next((p for p in lista if str(p.get('id')) == str(professor_id)), None)
    if not professor:
        flash('Professor não encontrado.', 'erro')
        return redirect(url_for('professores'))

    ano = int(request.form.get('ano', datetime.now().year))
    chart_b64 = request.form.get('chart_image', '')
    tipos_disponiveis = db.get_tipos_disponiveis_professor(professor_id, ano)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # Título
    titulo = doc.add_heading('', level=0)
    r_tit = titulo.add_run('DEVOLUTIVA DE ÍNDICES INDIVIDUAIS')
    r_tit.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_escola = doc.add_paragraph()
    r_esc = p_escola.add_run('E.E. Professora Dona Belinha — SEDUC-SP')
    r_esc.bold = True
    p_escola.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_dt = doc.add_paragraph(f'Gerado em {datetime.now().strftime("%d/%m/%Y")}')
    p_dt.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Info professor
    t_info = doc.add_table(rows=1, cols=3)
    t_info.style = 'Table Grid'
    for i, (label, val) in enumerate([
        ('Professor(a)', professor.get('nome', '')),
        ('CPF', professor.get('cpf') or 'Não informado'),
        ('Ano Letivo', str(ano)),
    ]):
        p = t_info.rows[0].cells[i].paragraphs[0]
        p.add_run(label + '\n').bold = True
        p.add_run(val)

    doc.add_paragraph()

    def _cell_shd(cell, fill_hex):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        tc_pr.append(shd)

    # Tabela por tipo
    for tipo in tipos_disponiveis:
        perf, bimestres = db.get_professor_performance(professor_id, ano, tipo)
        if not perf:
            continue

        h_tipo = doc.add_heading(tipo, level=2)
        if h_tipo.runs:
            h_tipo.runs[0].font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

        n_bim = len(bimestres)
        n_cols = 2 + n_bim + 3
        table = doc.add_table(rows=1, cols=n_cols)
        table.style = 'Table Grid'

        hdr_labels = (['Disciplina', 'Turma']
                      + [f'{b}° Bim' for b in bimestres]
                      + ['Méd. Prof.', 'Méd. Escola', 'Dif.'])
        for i, hl in enumerate(hdr_labels):
            cell = table.rows[0].cells[i]
            _cell_shd(cell, '1565C0')
            r = cell.paragraphs[0].add_run(hl)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i >= 2:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for item in perf:
            row = table.add_row().cells

            r0 = row[0].paragraphs[0].add_run(
                f"{DISCIPLINAS_NOMES.get(item['disciplina'], item['disciplina'])} ({item['disciplina']})")
            r0.font.size = Pt(9)

            r1 = row[1].paragraphs[0].add_run(item['turma'].split(' INTEGRAL')[0])
            r1.font.size = Pt(9)

            vals, medias_esc = [], []
            for j, b in enumerate(bimestres):
                dado = item['por_bimestre'].get(b)
                p_cell = row[2 + j].paragraphs[0]
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if dado and dado.get('valor'):
                    v = round(dado['valor'], 1)
                    r = p_cell.add_run(f'{v}%')
                    r.font.size = Pt(9)
                    r.font.color.rgb = (RGBColor(0x2E, 0x7D, 0x32) if v >= 50
                                        else RGBColor(0xC6, 0x28, 0x28))
                    vals.append(v)
                    if dado.get('media_escola'):
                        medias_esc.append(dado['media_escola'])
                else:
                    row[2 + j].paragraphs[0].add_run('—').font.size = Pt(9)

            mp_cell = row[2 + n_bim].paragraphs[0]
            mp_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            me_cell = row[2 + n_bim + 1].paragraphs[0]
            me_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dif_cell = row[2 + n_bim + 2].paragraphs[0]
            dif_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if vals:
                mp = round(sum(vals) / len(vals), 1)
                r_mp = mp_cell.add_run(f'{mp}%')
                r_mp.bold = True
                r_mp.font.size = Pt(9)
                r_mp.font.color.rgb = (RGBColor(0x2E, 0x7D, 0x32) if mp >= 50
                                        else RGBColor(0xC6, 0x28, 0x28))
                if medias_esc:
                    me = round(sum(medias_esc) / len(medias_esc), 1)
                    me_cell.add_run(f'{me}%').font.size = Pt(9)
                    dif = round(mp - me, 1)
                    r_dif = dif_cell.add_run(f'{"▲" if dif >= 0 else "▼"}{abs(dif)}')
                    r_dif.font.size = Pt(9)
                    r_dif.font.color.rgb = (RGBColor(0x2E, 0x7D, 0x32) if dif >= 0
                                            else RGBColor(0xC6, 0x28, 0x28))
                else:
                    me_cell.add_run('—').font.size = Pt(9)
                    dif_cell.add_run('—').font.size = Pt(9)
            else:
                mp_cell.add_run('—').font.size = Pt(9)
                me_cell.add_run('—').font.size = Pt(9)
                dif_cell.add_run('—').font.size = Pt(9)

        doc.add_paragraph()

    # Gráfico (imagem do canvas)
    if chart_b64 and 'base64,' in chart_b64:
        try:
            chart_bytes = _b64.b64decode(chart_b64.split('base64,')[1])
            h_graf = doc.add_heading('Evolução do Desempenho', level=2)
            if h_graf.runs:
                h_graf.runs[0].font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
            doc.add_picture(io.BytesIO(chart_bytes), width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e_g:
            print(f'Aviso: gráfico não incluído no Word: {e_g}')
        doc.add_paragraph()

    # Observações
    h_obs = doc.add_heading('Observações / Encaminhamentos:', level=3)
    if h_obs.runs:
        h_obs.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for _ in range(5):
        p_obs = doc.add_paragraph('_' * 85)
        p_obs.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()
    doc.add_paragraph()

    # Assinaturas
    t_sig = doc.add_table(rows=4, cols=2)
    for row_s in t_sig.rows:
        for c in row_s.cells:
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for col in range(2):
        t_sig.rows[0].cells[col].paragraphs[0].add_run('\n\n')

    for col in range(2):
        r = t_sig.rows[1].cells[col].paragraphs[0].add_run('_' * 38)
        r.bold = True

    t_sig.rows[2].cells[0].paragraphs[0].add_run('Diretor(a)').bold = True
    t_sig.rows[2].cells[1].paragraphs[0].add_run(professor.get('nome', 'Professor(a)')).bold = True
    t_sig.rows[3].cells[0].paragraphs[0].add_run('E.E. Professora Dona Belinha')
    t_sig.rows[3].cells[1].paragraphs[0].add_run('Professor(a)')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    nome = professor.get('nome', 'Professor').replace(' ', '_')
    return send_file(
        buf,
        as_attachment=True,
        download_name=f'Devolutiva_{nome}_{ano}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# ── Diagnóstico de CSV (rota temporária) ───────────────────────────────────

@app.route('/testar-csv', methods=['GET', 'POST'])
@login_required
def testar_csv():
    if request.method == 'POST':
        arquivo = request.files.get('arquivo')
        if not arquivo:
            return jsonify({'erro': 'Nenhum arquivo enviado'})
        try:
            conteudo = arquivo.read()
            # Detecta encoding
            encoding_usado = None
            for enc in ('utf-8-sig', 'utf-8', 'cp1252', 'latin-1'):
                try:
                    conteudo.decode(enc)
                    encoding_usado = enc
                    break
                except Exception:
                    continue

            registros = csv_parser.parse_csv(conteudo)
            return jsonify({
                'encoding_detectado': encoding_usado,
                'total_registros': len(registros),
                'primeiros_3': registros[:3],
            })
        except Exception as e:
            return jsonify({'erro': str(e)})

    return '''
    <form method="POST" enctype="multipart/form-data">
        <input type="file" name="arquivo" accept=".csv"><br><br>
        <button type="submit">Testar CSV</button>
    </form>
    '''


# ── Diagnóstico do banco de dados ─────────────────────────────────────────

@app.route('/testar-db')
@login_required
def testar_db():
    conn = db.get_db()
    tipos = conn.execute(
        'SELECT tipo_avaliacao, COUNT(*) as qtd FROM avaliacoes GROUP BY tipo_avaliacao'
    ).fetchall()
    vinculos_lista = conn.execute(
        'SELECT professor_id, nome_professor, disciplina, turma, ano_letivo FROM vinculos ORDER BY nome_professor'
    ).fetchall()
    # Para cada vínculo, mostra se existe avaliação SARESP com aquela turma
    detalhes = []
    for v in vinculos_lista:
        av = conn.execute(
            "SELECT tipo_avaliacao, COUNT(*) as qtd FROM avaliacoes WHERE turma=? AND ano=? GROUP BY tipo_avaliacao",
            (v['turma'], v['ano_letivo'])
        ).fetchall()
        detalhes.append({
            'professor': v['nome_professor'],
            'disciplina': v['disciplina'],
            'turma': v['turma'],
            'ano_letivo': v['ano_letivo'],
            'avaliacoes': [dict(a) for a in av],
        })
    conn.close()
    return jsonify({
        'tipos_no_banco': [dict(t) for t in tipos],
        'vinculos_com_avaliacoes': detalhes,
    })


# ── Sincronização cache ────────────────────────────────────────────────────

def _sincronizar_cache():
    """Sincroniza vínculos e avaliações da planilha → SQLite via API (tempo real, sem cache)."""
    resultado = {'vinculos': 0, 'avaliacoes': 0, 'erros': []}
    try:
        rows = sheets.get_vinculos_planilha()
        resultado['vinculos'] = db.sync_vinculos(rows)
    except Exception as e:
        resultado['erros'].append(f'vínculos: {e}')
    try:
        rows = sheets.get_avaliacoes_planilha()
        resultado['avaliacoes'] = db.sync_avaliacoes(rows)
    except Exception as e:
        resultado['erros'].append(f'avaliações: {e}')
    return resultado


# ── Inicialização ──────────────────────────────────────────────────────────

if __name__ == '__main__':
    db.init_db()
    print('\n  Sistema Belinha iniciado!')
    print('  Sincronizando com Google Sheets (API)...')
    r = _sincronizar_cache()
    print(f'  Cache: {r["vinculos"]} vínculos · {r["avaliacoes"]} avaliações carregadas')
    if r['erros']:
        for e in r['erros']:
            print(f'  ERRO: {e}')
    print('  Acesse: http://127.0.0.1:5000\n')
    app.run(debug=True)
